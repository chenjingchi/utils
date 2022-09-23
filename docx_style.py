import docx
import os
import re

from docx.shared import Pt


def txt_to_docx(txt):
    """Converts text to docx with font style adjusted.

    Tibetan: font size: 20, name: 'Microsoft Himalaya'
    Chinese: font size 10.5 (“小四"), name: '仿宋'

    Args:
        txt: str, input text.

    Returns:
        docx object with the font adjusted.
    """

    tibetan_font_size = 20
    tibetan_font_name = 'Microsoft Himalaya'

    chinese_font_size = 10.5  # 小四
    chinese_font_name = '仿宋'

    doc = docx.Document()
    for idx, para_txt in enumerate(txt.split("\n")):
        # Skip digit line.
        if re.match(r'[\.0-9]+', para_txt):
            continue
        para = doc.add_paragraph()
        font = para.add_run(para_txt).font
        # is Tibetan
        if re.findall(r'[\u0f00-\u0fff]+', para_txt):
            font.size = Pt(tibetan_font_size)
            font.name = tibetan_font_name
        else:
            font.size = Pt(chinese_font_size)
            font.name = chinese_font_name
    return doc


INPUT_ROOT = "path/to/input/txt/directory"
OUTPUT_ROOT = "path/to/output/docx/directory"

for root, dirs, files in os.walk(INPUT_ROOT):
    for filename in files:
        print(filename)
        if filename.endswith(".txt"):
            fullpath = os.path.join(root, filename)
            print(fullpath)
            with open(fullpath, "rb") as file:
                txt = file.read().decode('utf-8')

                doc = txt_to_docx(txt)

                output_filename = os.path.splitext(filename)[0]
                output_dir = os.path.join(OUTPUT_ROOT, os.path.basename(root))
                if not os.path.exists(output_dir):
                    os.makedirs(output_dir)
                output_fullpath = os.path.join(
                    output_dir, output_filename) + ".docx"
                doc.save(output_fullpath)
